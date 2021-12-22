from docx import Document
from docx.enum.table import WD_TABLE_DIRECTION
import sys
import os

cfactor_path = os.path.abspath(os.path.join(os.path.dirname(__file__), os.path.pardir))


def export(building=None, filename=None):
    if not building:
        return None

    X = building.x_system
    Y = building.y_system

    prop = {'': '',
            'محل اجرای پروژه': building.city,
            'کاربری ساختمان': 'مسکونی',
            'ضریب اهمیت': building.importance_factor,
            'تعداد طبقات': building.number_of_story,
            'ارتفاع ساختمان  )متر(': building.height,
            'سطح خطر نسبی': building.risk_level,
            'شتاب مبنای طرح': building.acc,
            'نوع خاک': building.soilType}
    if X == Y:
        prop['سیستم سازه ای در دو راستا'] = X.lateralType
    else:
        prop['سیستم سازه ای در راستای x'] = X.lateralType
        prop['سیستم سازه ای در راستای y'] = Y.lateralType

    struc = {'': ('راستای x', 'راستای y'),
             'سیستم سازه': (X.lateralType, Y.lateralType),
             'ضریب رفتار': (X.Ru, Y.Ru),
             'ضریب اضافه مقاومت': (X.phi0, Y.phi0),
             'ضریب بزرگنمایی جابجایی': (X.cd, Y.cd),
             'ارتفاع مجاز  )متر(': (X.maxHeight, Y.maxHeight)}

    result = {'زمان تناوب تجربی': (building.exp_period_x, building.exp_period_y),
              'زمان تناوب تحلیلی': (building.x_period_an, building.y_period_an),
              'ضریب بازتاب': (building.Bx, building.By),
              'C': (building.results[1], building.results[2]),
              'K': (building.kx, building.ky),
              'C_drift': (building.results_drift[1], building.results_drift[2]),
              'K_drift': (building.kx_drift, building.ky_drift),
              }

    doc = Document(os.path.join(cfactor_path, 'exporter', 'template.docx'))
    doc.add_heading('محاسبه ضریب زلزله', level=0)
    doc.add_heading('مشخصات پروژه', level=1)
    table_prop = doc.add_table(rows=0, cols=2, style=doc.styles['List Table 4 Accent 5'])
    table_prop.direction = WD_TABLE_DIRECTION.RTL
    for key, value in prop.items():
        row_cells = table_prop.add_row().cells
        row_cells[0].text = key
        row_cells[1].text = str(value)

    doc.add_heading('مشخصات سیستم سازه ای', level=1)
    struc_table = doc.add_table(rows=0, cols=3, style=doc.styles['List Table 4 Accent 5'])
    for key, value in struc.items():
        row_cells = struc_table.add_row().cells
        row_cells[0].text = key
        row_cells[1].text = str(value[0])
        row_cells[2].text = str(value[1])

    doc.add_heading('ضریب زلزله', level=1)
    result_table = doc.add_table(rows=1, cols=3, style=doc.styles['List Table 4 Accent 5'])
    hdr_cells = result_table.rows[0].cells
    hdr_cells[1].text = 'راستای x'
    hdr_cells[2].text = 'راستای y'
    for key, value in result.items():
        row_cells = result_table.add_row().cells
        row_cells[0].text = key
        row_cells[1].text = f'{value[0]:.3f}'
        row_cells[2].text = f'{value[1]:.3f}'

    if filename:
        doc.save(filename)
