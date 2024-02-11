import sys
from pathlib import Path
from decimal import Decimal
import math

import numpy as np

import matplotlib.pyplot as plt
import matplotlib.patches as patches
try:
    from docx.shared import Inches
except ImportError:
    import subprocess
    import sys
    package = 'python-docx'
    subprocess.check_call(['python', "-m", "pip", "install", package])


from python_functions import get_temp_filepath
from osafe_funcs import  osafe_funcs
import Part

import FreeCAD
import Part

def add_edges_to_ax(obj, ax, linewidth=.5, linestyle='-'):
    for e in obj.Edges:
        if len(e.Vertexes) != 2:
            continue
        v1, v2 = e.Vertexes
        xy = [[v1.X, v1.Y], [v2.X, v2.Y]]
        p = patches.Polygon(xy, edgecolor='black', linewidth=linewidth, 
                            linestyle=linestyle, closed=False,
                            )
        ax.add_patch(p)

def add_key_plan_edges_to_ax(punch, ax):
    comp = Part.makeCompound([
            punch.foundation.plan,
            punch.edges,
        ])
    add_edges_to_ax(comp, ax, linewidth=.1)
    add_column_to_ax(punch, ax)
    set_ax_boundbox(comp, ax)
    return comp

def add_column_to_ax(punch, ax):
    if FreeCAD.GuiUp:
        color = punch.ViewObject.ShapeColor[:-1]
    else:
        color = (0.5, 0.5, 0.5)
    z_min = punch.column.Shape.BoundBox.ZMin
    for face in punch.column.Shape.Faces:
        if face.BoundBox.ZMin == z_min and face.BoundBox.ZMax == z_min:
            break
    xy = [[v.X, v.Y] for v in face.OuterWire.Vertexes]
    xy = osafe_funcs.sort_vertex(xy)
    p = patches.Polygon(xy, edgecolor='black', linewidth=.8, facecolor=color, closed=True)
    ax.add_patch(p)

def add_base_plate_to_ax(punch, ax):
    if hasattr(punch.column, 'base_plate') and punch.column.base_plate:
        base_plate = punch.column.base_plate
        color = (1, 0 , 0)
        if FreeCAD.GuiUp:
            color = base_plate.ViewObject.ShapeColor[:-1]
    else:
        return
    # Draw base plate
    z_min = base_plate.Shape.BoundBox.ZMin
    for face in base_plate.Shape.Faces:
        if face.BoundBox.ZMin == z_min and face.BoundBox.ZMax == z_min:
            break
    xy = [[v.X, v.Y] for v in face.Vertexes]
    xy = osafe_funcs.sort_vertex(xy)
    p = patches.Polygon(xy, edgecolor='black', linewidth=.8, facecolor=color, closed=True)
    ax.add_patch(p)
    # Add equivalent column
    xy = [[v.X, v.Y] for v in punch.rect.Vertexes]
    xy = osafe_funcs.sort_vertex(xy)
    p = patches.Polygon(xy, edgecolor='black', linewidth=.6, closed=False)
    ax.add_patch(p)

def add_punch_edges_to_ax(punch, ax):
    edges = punch.edges
    add_edges_to_ax(edges, ax, linewidth=2)
    ret_null_edges, common_edges = osafe_funcs.punch_null_edges(punch=punch)
    null_edges = []
    for ret, edge in zip(ret_null_edges, common_edges):
        if ret == 'Yes':
            null_edges.append(edge)
    null_edges = Part.makeCompound(null_edges)
    add_edges_to_ax(null_edges, ax, linewidth=.1)
    add_base_plate_to_ax(punch, ax)
    add_column_to_ax(punch, ax)
    set_ax_boundbox(common_edges, ax)
    # return comp

def get_column_points_for_dimensioning(col):
    z_min = col.Shape.BoundBox.ZMin
    for face in col.Shape.Faces:
        if face.BoundBox.ZMin == z_min and face.BoundBox.ZMax == z_min:
            break
    e1, e2 = Part.__sortEdges__(face.OuterWire.Edges)[2:]
    v1, v2 = e1.Vertexes
    v3, v4 = e2.Vertexes
    return v1, v2, v3, v4

def add_punch_edges_dimension_to_ax(punch, ax):
    v1, v2, v3, v4 = get_column_points_for_dimensioning(punch.column)
    offset = 60
    annotate_dim(ax,(v1.X, v1.Y),(v2.X, v2.Y), offset=offset, arrows=False)
    annotate_dim(ax,(v3.X, v3.Y),(v4.X, v4.Y), offset=offset, arrows=False)
    for e in punch.edges.Edges:
        v1, v2 = e.Vertexes
        annotate_dim(ax,(v1.X, v1.Y), (v2.X, v2.Y), offset=-offset)

def set_ax_boundbox(shape, ax, scale=1, bb_add=120):
    if isinstance(shape, list):
        shape = Part.makeCompound(shape)
    b = shape.BoundBox
    xmin, ymin, xmax, ymax = b.XMin, b.YMin, b.XMax, b.YMax
    ax.set_xlim((xmin - bb_add) * scale, (xmax + bb_add) * scale)
    ax.set_ylim((ymin - bb_add) * scale, (ymax + bb_add) * scale)

def get_punch_picture(
        punch,
        filename='',
        ) -> Path:
    # fig, axes = plt.subplots(2, 1)
    # for ax in axes:
    #     ax.set_aspect('equal')
    #     ax.axis('off')
    # ax1, ax2 = axes
    fig = plt.figure(constrained_layout=True)
    widths = [2, 2, 2]
    heights = [2, 2, 1]
    gs = fig.add_gridspec(ncols=3, nrows=3, width_ratios=widths,
                          height_ratios=heights, left=0.05, right=0.48, wspace=0)
    ax1 = fig.add_subplot(gs[:-1, :], aspect='equal')
    plt.axis('off')
    add_punch_edges_to_ax(punch, ax1)
    add_punch_edges_dimension_to_ax(punch, ax1)
    # ax2 = fig.add_subplot(2, 1, 2, aspect='equal')
    ax2 = fig.add_subplot(gs[-1, 1], aspect='equal')
    plt.axis('off')
    add_key_plan_edges_to_ax(punch, ax2)
    if not filename or not Path(filename).exists():
        filename = get_temp_filepath(suffix='jpg', filename=punch.Name)
    fig.savefig(str(filename), orientation='portrait',
        # papertype='a4',
        bbox_inches='tight', dpi=600)
    plt.clf()
    plt.close(fig)
    return filename

def move_xy(xy, dist=200, angle=0):
    # return xy
    x_offset = dist * np.cos(angle + np.pi/2)
    y_offset = dist * np.sin(angle + np.pi/2)
    return (xy[0] + x_offset, xy[1] + y_offset)

def annotate_dim(ax,xyfrom,xyto, offset=100, text=None, arrows=True):
    '''
    direction: dimension arrow offset. can be: 'TOP', 'BOT', 'LEFT', 'RIGHT'
    '''
    delta_x = xyto[0] - xyfrom[0]
    delta_y = xyto[1] - xyfrom[1]
    angle = np.arctan2(delta_y, delta_x)
    fontsize = 5
    if text is None:
        dist = np.sqrt((xyfrom[0] - xyto[0]) ** 2 + (xyfrom[1] - xyto[1]) ** 2)
        text = f'{dist / 10: .0f} cm'
        if not arrows and dist < 300:
            offset = 20
        #     fontsize = 3
    xyfrom = move_xy(xyfrom, offset, angle=angle)
    xyto = move_xy(xyto, offset, angle=angle)
    if abs(angle) > math.pi / 2:
        angle = angle - math.pi
    color = 'blue'
    if arrows:
        ax.annotate("",xyfrom,xyto, arrowprops=dict(arrowstyle='<->', lw=.2, color=color))
    # ax.annotate("",xyfrom,xyto,arrowprops=dict(arrowstyle='|-|', lw=.2, edgecolor=color))
    bbox=dict(fc="white", ec="none", pad=0.1)
    ax.text((xyto[0]+xyfrom[0])/2,(xyto[1]+xyfrom[1])/2,text,fontsize=fontsize,
        color=color,ha="center", va="center", bbox=bbox,
        rotation=np.degrees(angle),
        )

def export_dataframe_to_docx(df, doc=None):
    if not doc:
        doc = create_doc()

    # add a table to the end and create a reference variable
    # extra row is so we can add the header row
    t = doc.add_table(df.shape[0]+1, df.shape[1], style='DATAFRAME')
    from docx.enum.table import WD_TABLE_ALIGNMENT
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    # add the header rows.
    for j in range(df.shape[-1]):
        cell = t.cell(0, j)
        cell.text = df.columns[j]
        # run = cell.paragraphs[0].runs[0]
        # run.font.bold = True

    # add the rest of the data frame
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            t.cell(i+1,j).text = str(df.values[i,j])
    return doc

def export_dict_to_doc(
                        d: dict,
                        headers: list,
                        doc: 'docx.Document' = None,
                        style: str = 'vazhgooni'
                        ) -> 'docx.Document':
    if not doc:
        doc = create_doc()
    num_rows = len(d)
    num_cols = 2
    start_row = 0
    if headers:
        num_rows += 1
        start_row = 1
    # add a table to the end and create a reference variable
    # extra row is so we can add the header row
    t = doc.add_table(num_rows, num_cols, style=style)
    from docx.enum.table import WD_TABLE_ALIGNMENT
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    # add the header rows.
    for j, header in enumerate(headers):
        cell = t.cell(0, j)
        cell.text = str(header)
        # run = cell.paragK

    # add the rest of the data frame

    for i, (key, value) in enumerate(d.items(), start=start_row):
        t.cell(i,0).text = str(key)
        t.cell(i,1).text = str(value)
    return doc
    

def add_punch_ratios_to_doc(punch, doc=None):
    doc = export_dict_to_doc(punch.combos_ratio, ['Combo', 'Ratio'], doc)
    return doc

def add_punch_properties_to_doc(punch, doc=None):
    if not doc:
        doc = create_doc()
    d = dict()
    d['Punch ID'] = punch.id
    d['Ratio'] = punch.Ratio
    d['location'] = punch.Location
    d['X Coordinate'] = f'{punch.center_of_load.x} mm'
    d['Y Coordinate'] = f'{punch.center_of_load.y} mm'
    d['perimeter'] = f'{punch.b0:.0f} mm'
    d['Eff. depth'] = f'{punch.d:.0f} mm'
    d['Cover'] = f'{punch.foundation.cover.Value:.0f} mm'
    d['I22'] = f'{Decimal(punch.I22):.2E} mm^4'
    d['I33'] = f'{Decimal(punch.I33):.2E} mm^4'
    d['I23'] = f'{Decimal(punch.I23):.2E} mm^4'
    d['fc'] = f'{float(punch.fc.getValueAs("N/mm^2")):.1f} Mpa'
    d['gamma vx'] = f'{punch.gamma_vx:.2f}'
    d['gamma vy'] = f'{punch.gamma_vy:.2f}'
    doc = export_dict_to_doc(d, [], doc, 'connection')
    return doc

def create_doc():
    import docx
    punch_path = Path(__file__).parent.parent
    filepath = punch_path / 'osafe_import_export' / 'templates' / 'punch_default.docx'
    doc = docx.Document(str(filepath))
    return doc

def create_report(
                punch,
                filename: str = None,
                doc: 'docx.Document' = None,
                ):
    if not doc:
        doc = create_doc()
    doc.add_heading(f'PUNCH ID = {punch.id}',0)
    image_file_path = get_punch_picture(punch)
    doc.add_picture(str(image_file_path), width=Inches(3.5))
    doc.add_paragraph()
    doc = add_punch_properties_to_doc(punch, doc)
    doc.add_paragraph()
    doc = add_punch_ratios_to_doc(punch, doc)
    if filename:
        doc.save(filename)
    return doc

def create_punches_report(
                    document,
                    filename: str,
                    doc: 'docx.Document' = None,
    ):
    if not doc:
        doc = create_doc()
    punches = []
    for o in document.Objects:
        if hasattr(o, 'Proxy') and hasattr(o.Proxy, 'Type') and o.Proxy.Type == 'Punch':
            punches.append(o)
    progressbar = FreeCAD.Base.ProgressIndicator()
    n = len(punches)
    progressbar.start("export "+str(n)+" Punches to Word Document ...", n)
    for p in punches:
        doc = create_report(p, doc=doc)
        doc.add_page_break()
        progressbar.next(True)
    if filename:
        doc.save(filename)
    progressbar.stop()
    return doc



