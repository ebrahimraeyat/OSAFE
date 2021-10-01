import sys
from pathlib import Path
from decimal import Decimal

import numpy as np

import matplotlib.pyplot as plt
import matplotlib.patches as patches
try:
	from docx.shared import Inches
except ImportError:
	import subprocess
	import sys
	package = 'python-docx'
	subprocess.check_call(['python', "-m", "pip", "install", package])



punch_path = Path(__file__).absolute().parent
try:
	import safe.punch.punch_funcs as punch_funcs
except:
	import  punch_funcs
import Part

try:
	import FreeCAD
except:
	# path to FreeCAD.so
	FREECADPATH = '/usr/lib/freecad-daily-python3/lib/'
	sys.path.append(FREECADPATH)
	import FreeCAD


def add_edges_to_ax(obj, ax, linewidth=.5):
	for e in obj.Edges:
		if not len(e.Vertexes) == 2:
			continue
		v1, v2 = e.Vertexes
		xy = [[v1.X, v1.Y], [v2.X, v2.Y]]
		p = patches.Polygon(xy, edgecolor='black', linewidth=linewidth, closed=False)
		ax.add_patch(p)


def add_key_plan_edges_to_ax(punch, ax):
	comp = Part.makeCompound([
			punch.foundation.plane,
			punch.edges,
		])
	add_edges_to_ax(comp, ax, linewidth=.1)
	add_column_to_ax(punch, ax)
	set_ax_boundbox(comp, ax)
	return comp

def add_column_to_ax(punch, ax):
	if FreeCAD.GuiUp:
		color = punch.ViewObject.ShapeColor[:-1]
	else:
		color = (.5, 0.5, 0.5)
	xy = [[v.X, v.Y] for v in punch.rect.Vertexes]
	xy = punch_funcs.sort_vertex(xy)
	p = patches.Polygon(xy, edgecolor='black', linewidth=.8, facecolor=color, closed=True)
	ax.add_patch(p)


def add_punch_edges_to_ax(punch, ax):
	edges = punch.edges
	add_edges_to_ax(edges, ax, linewidth=.8)
	add_column_to_ax(punch, ax)
	set_ax_boundbox(edges, ax)
	# return comp

def add_punch_edges_dimension_to_ax(punch, ax):
	xmin = punch.center_of_load.x - punch.bx / 2
	ymin = punch.center_of_load.y - punch.by / 2
	x2 = xmin + punch.bx
	y2 = ymin + punch.by
	annotate_dim(ax,(xmin, ymin),(x2, ymin), direction='BOT', offset=80)
	annotate_dim(ax,(xmin, ymin),(xmin, y2), direction='LEFT', offset=100)
	edges_dim = get_edges_direction_in_punch(punch)
	for e, loc in edges_dim.items():
		v1, v2 = e.Vertexes
		if loc in ('LEFT', 'RIGHT'):
			offset = 120
		else:
			offset = 80
		annotate_dim(ax,(v1.X, v1.Y), (v2.X, v2.Y), direction=loc, offset=offset,
			rotation=48)

def set_ax_boundbox(shape, ax, scale=1, bb_add=120):
	b = shape.BoundBox
	xmin, ymin, xmax, ymax = b.XMin, b.YMin, b.XMax, b.YMax
	ax.set_xlim((xmin - bb_add) * scale, (xmax + bb_add) * scale)
	ax.set_ylim((ymin - bb_add) * scale, (ymax + bb_add) * scale)

def get_edges_direction_in_punch(punch):
	edges_direction = {}
	x_center = punch.center_of_load.x
	y_center = punch.center_of_load.y
	for e in punch.edges.Edges:
		bb = e.BoundBox
		if bb.XLength <= bb.YLength:
			if bb.XMax > x_center:
				edges_direction[e] = 'RIGHT'
			elif bb.XMax < x_center:
				edges_direction[e] = 'LEFT'
		elif bb.XLength > bb.YLength:
			if bb.YMax > y_center:
				edges_direction[e] = 'TOP'
			elif bb.YMax < y_center:
				edges_direction[e] = 'BOT'
	return edges_direction

def get_punch_picture(
		punch,
		filename='',
		) -> Path:
	fig = plt.figure(constrained_layout=True)
	widths = [2, 2, 2]
	heights = [2, 2, 1]
	gs = fig.add_gridspec(ncols=3, nrows=3, width_ratios=widths,
                          height_ratios=heights, left=0.05, right=0.48, wspace=0)
	ax1 = fig.add_subplot(gs[:-1, :], aspect='equal')
	plt.axis('off')
	add_punch_edges_to_ax(punch, ax1)
	add_punch_edges_dimension_to_ax(punch, ax1)
	ax2 = fig.add_subplot(gs[-1, 1], aspect='equal')
	plt.axis('off')
	add_key_plan_edges_to_ax(punch, ax2)
	import tempfile
	default_tmp_dir = tempfile._get_default_tempdir()
	if not filename:
		filename = punch.Name
	filename_path = Path(default_tmp_dir) / f'{filename}.jpg'
	fig.savefig(str(filename_path), orientation='portrait',
		# papertype='a4',
		bbox_inches='tight', dpi=600)
	plt.clf()
	plt.close(fig)
	return filename_path

def move_xy(xy, direction, dist=200):
	if direction in ('BOT', 'LEFT'):
		dist = -dist
	if direction in  ('TOP', 'BOT'):
		return (xy[0], xy[1] + dist)
	elif direction in ('LEFT', 'RIGHT'):
		return (xy[0] + dist, xy[1])


def annotate_dim(ax,xyfrom,xyto, direction='TOP', offset=200, text=None, rotation=0):
	'''
	direction: dimension arrow offset. can be: 'TOP', 'BOT', 'LEFT', 'RIGHT'
	'''
	xyfrom = move_xy(xyfrom, direction, offset)
	xyto = move_xy(xyto, direction, offset)
	if text is None:
		dist = np.sqrt((xyfrom[0]-xyto[0])**2 + (xyfrom[1]-xyto[1])**2)
		text = f'{dist: .0f}'
	color = 'grey'
	ax.annotate("",xyfrom,xyto,arrowprops=dict(arrowstyle='<->', lw=.2, edgecolor=color))
	ax.annotate("",xyfrom,xyto,arrowprops=dict(arrowstyle='|-|', lw=.2, edgecolor=color))
	bbox=dict(fc="white", ec="none")
	ax.text((xyto[0]+xyfrom[0])/2,(xyto[1]+xyfrom[1])/2,text,fontsize=7,
		color=color,ha="center", va="center", bbox=bbox,
		rotation=rotation)

def export_dataframe_to_docx(df, doc=None):
	if not doc:
		doc = create_doc()

	# add a table to the end and create a reference variable
	# extra row is so we can add the header row
	t = doc.add_table(df.shape[0]+1, df.shape[1], style='DATAFRAME')
	from docx.enum.table import WD_TABLE_ALIGNMENT
	t.alignment = WD_TABLE_ALIGNMENT.CENTER

	# add the header rows.
	for j in range(df.shape[-1]):
		cell = t.cell(0, j)
		cell.text = df.columns[j]
		# run = cell.paragraphs[0].runs[0]
		# run.font.bold = True

	# add the rest of the data frame
	for i in range(df.shape[0]):
		for j in range(df.shape[-1]):
			t.cell(i+1,j).text = str(df.values[i,j])
	return doc

def export_dict_to_doc(
						d: dict,
						headers: list,
						doc: 'docx.Document' = None,
						style: str = 'vazhgooni'
						) -> 'docx.Document':
	if not doc:
		doc = create_doc()
	num_rows = len(d)
	num_cols = 2
	start_row = 0
	if headers:
		num_rows += 1
		start_row = 1
	# add a table to the end and create a reference variable
	# extra row is so we can add the header row
	t = doc.add_table(num_rows, num_cols, style=style)
	from docx.enum.table import WD_TABLE_ALIGNMENT
	t.alignment = WD_TABLE_ALIGNMENT.CENTER

	# add the header rows.
	for j, header in enumerate(headers):
		cell = t.cell(0, j)
		cell.text = str(header)
		# run = cell.paragK

	# add the rest of the data frame

	for i, (key, value) in enumerate(d.items(), start=start_row):
		t.cell(i,0).text = str(key)
		t.cell(i,1).text = str(value)
	return doc
	

def add_punch_ratios_to_doc(punch, doc=None):
	doc = export_dict_to_doc(punch.combos_ratio, ['Combo', 'Ratio'], doc)
	return doc

def add_punch_properties_to_doc(punch, doc=None):
	if not doc:
		doc = create_doc()
	d = dict()
	d['Punch ID'] = punch.id
	d['Ratio'] = punch.Ratio
	d['location'] = punch.Location
	d['X Coordinate'] = f'{punch.center_of_load.x} mm'
	d['Y Coordinate'] = f'{punch.center_of_load.y} mm'
	d['perimeter'] = f'{punch.b0:.0f} mm'
	d['Eff. depth'] = f'{punch.d:.0f} mm'
	d['Cover'] = f'{punch.foundation.cover.Value:.0f} mm'
	d['I22'] = f'{Decimal(punch.I22):.2E} mm^4'
	d['I33'] = f'{Decimal(punch.I33):.2E} mm^4'
	d['I23'] = f'{Decimal(punch.I23):.2E} mm^4'
	d['fc'] = f'{float(punch.fc.getValueAs("N/mm^2")):.1f} Mpa'
	d['gamma vx'] = f'{punch.gamma_vx:.2f}'
	d['gamma vy'] = f'{punch.gamma_vy:.2f}'
	doc = export_dict_to_doc(d, [], doc, 'connection')
	return doc

def create_doc():
	import docx
	filepath = punch_path / 'Resources' / 'templates' / 'default.docx'
	doc = docx.Document(str(filepath))
	return doc

def create_report(
				punch,
				filename: str = None,
				doc: 'docx.Document' = None,
				):
	if not doc:
		doc = create_doc()
	doc.add_heading(f'PUNCH ID = {punch.id}',0)
	image_file_path = get_punch_picture(punch)
	doc.add_picture(str(image_file_path), width=Inches(3.5))
	doc.add_paragraph()
	doc = add_punch_properties_to_doc(punch, doc)
	doc.add_paragraph()
	doc = add_punch_ratios_to_doc(punch, doc)
	if filename:
		doc.save(filename)
	return doc

def create_punches_report(
					document,
					filename: str,
					doc: 'docx.Document' = None,
):
	if not doc:
		doc = create_doc()
	for o in document.Objects:
		if 'Punch' in o.Name:
			doc = create_report(o, doc=doc)
			doc.add_page_break()
	if filename:
		doc.save(filename)
	return doc



